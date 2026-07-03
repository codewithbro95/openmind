from __future__ import annotations

from pathlib import Path

from openmind.core.models import ExtractedDocument
from openmind.extractors.base import Extractor


class DocxExtractor(Extractor):
    extensions = {".docx"}

    def extract(self, path: str) -> ExtractedDocument:
        from docx import Document

        file_path = Path(path)
        document = Document(str(file_path))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return ExtractedDocument(
            file_path=str(file_path),
            title=file_path.stem,
            text="\n".join(paragraphs),
            metadata={"extension": ".docx"},
        )
